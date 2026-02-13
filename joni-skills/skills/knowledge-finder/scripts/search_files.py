#!/usr/bin/env python3
"""
Knowledge Finder - Exhaustive Research File Search & Extraction

Searches research files for relevant content with precise location tracking.
Returns ALL matches with exact source attribution for bulletproof traceability.

Supports: .txt, .md, .docx, .pdf, .xlsx, .csv, .json, .html

Usage:
    python search_files.py <folders> --query "search terms" [--output results.json]
    python search_files.py ./research ./interviews --query "onboarding friction"
"""

import argparse
import json
import os
import re
import sys
from pathlib import Path
from dataclasses import dataclass, field, asdict
from typing import Optional
from datetime import datetime
import hashlib


@dataclass
class Evidence:
    """A single piece of evidence with full attribution."""
    content: str                    # The actual quote or data point
    file_path: str                  # Full path to source file
    file_name: str                  # Just the filename
    file_type: str                  # Extension
    location: str                   # Precise location (page, row, timestamp, etc.)
    location_type: str              # Type of location marker (page, row, paragraph, timestamp, slide)
    context_before: str             # Text immediately before the match
    context_after: str              # Text immediately after the match
    speaker: Optional[str] = None   # Speaker/participant if applicable
    date_detected: Optional[str] = None  # Date found in file or metadata
    match_score: float = 0.0        # Relevance score
    search_terms_matched: list = field(default_factory=list)  # Which query terms matched


@dataclass 
class FileAnalysis:
    """Analysis results for a single file."""
    file_path: str
    file_name: str
    file_type: str
    file_size: int
    modified_date: str
    evidence_count: int
    evidence: list[Evidence] = field(default_factory=list)
    extraction_errors: list[str] = field(default_factory=list)
    is_research_file: bool = True   # Flag if file appears to be research-related


@dataclass
class SearchResults:
    """Complete search results across all files."""
    query: str
    search_terms: list[str]
    folders_searched: list[str]
    files_analyzed: int
    files_with_matches: int
    total_evidence_count: int
    search_timestamp: str
    file_analyses: list[FileAnalysis] = field(default_factory=list)
    errors: list[str] = field(default_factory=list)


# ============================================================================
# TEXT EXTRACTION BY FILE TYPE
# ============================================================================

def extract_text_from_txt(file_path: str) -> list[dict]:
    """Extract text from plain text/markdown files with line numbers."""
    segments = []
    try:
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            lines = f.readlines()
        
        current_speaker = None
        speaker_patterns = [
            r'^(?:Speaker\s*\d+|P\d+|Participant\s*\d+|Interviewer|Moderator|I:|P:)\s*[:\-]?\s*',
            r'^([A-Z][a-z]+(?:\s+[A-Z]\.?)?)\s*:\s*',  # "John:" or "John S.:"
            r'^\[([^\]]+)\]\s*:?\s*',  # [Speaker Name]
            r'^\*\*([^*]+)\*\*\s*:?\s*',  # **Speaker**:
        ]
        
        for i, line in enumerate(lines, 1):
            line_stripped = line.strip()
            if not line_stripped:
                continue
            
            # Check for speaker labels
            detected_speaker = None
            for pattern in speaker_patterns:
                match = re.match(pattern, line_stripped, re.IGNORECASE)
                if match:
                    detected_speaker = match.group(1) if match.lastindex else match.group(0).strip(':- ')
                    current_speaker = detected_speaker
                    break
            
            segments.append({
                'content': line_stripped,
                'location': f'Line {i}',
                'location_type': 'line',
                'line_number': i,
                'speaker': current_speaker,
                'context_before': lines[i-2].strip() if i > 1 else '',
                'context_after': lines[i].strip() if i < len(lines) else '',
            })
    except Exception as e:
        return [{'error': str(e)}]
    
    return segments


def extract_text_from_docx(file_path: str) -> list[dict]:
    """Extract text from Word documents with paragraph tracking."""
    segments = []
    try:
        from docx import Document
        doc = Document(file_path)
        
        para_num = 0
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            para_num += 1
            
            # Check for speaker patterns in paragraphs
            speaker = None
            speaker_match = re.match(r'^(?:\*\*)?([^:*]+)(?:\*\*)?\s*:\s*', text)
            if speaker_match:
                speaker = speaker_match.group(1).strip()
            
            segments.append({
                'content': text,
                'location': f'Paragraph {para_num}',
                'location_type': 'paragraph',
                'paragraph_number': para_num,
                'speaker': speaker,
                'style': para.style.name if para.style else None,
                'context_before': '',
                'context_after': '',
            })
        
        # Also extract from tables
        table_num = 0
        for table in doc.tables:
            table_num += 1
            for row_idx, row in enumerate(table.rows, 1):
                for col_idx, cell in enumerate(row.cells, 1):
                    text = cell.text.strip()
                    if text:
                        segments.append({
                            'content': text,
                            'location': f'Table {table_num}, Row {row_idx}, Col {col_idx}',
                            'location_type': 'table_cell',
                            'table_number': table_num,
                            'row': row_idx,
                            'column': col_idx,
                            'speaker': None,
                            'context_before': '',
                            'context_after': '',
                        })
        
        # Add context
        for i, seg in enumerate(segments):
            if i > 0:
                seg['context_before'] = segments[i-1]['content'][:100]
            if i < len(segments) - 1:
                seg['context_after'] = segments[i+1]['content'][:100]
                
    except ImportError:
        return [{'error': 'python-docx not installed. Run: pip install python-docx'}]
    except Exception as e:
        return [{'error': str(e)}]
    
    return segments


def extract_text_from_pdf(file_path: str) -> list[dict]:
    """Extract text from PDFs with page tracking."""
    segments = []
    try:
        import pdfplumber
        
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                text = page.extract_text()
                if not text:
                    continue
                
                # Split into paragraphs
                paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
                for para_idx, para in enumerate(paragraphs, 1):
                    segments.append({
                        'content': para,
                        'location': f'Page {page_num}, Paragraph {para_idx}',
                        'location_type': 'page',
                        'page_number': page_num,
                        'paragraph_in_page': para_idx,
                        'speaker': None,
                        'context_before': '',
                        'context_after': '',
                    })
                
                # Also extract tables
                tables = page.extract_tables()
                for table_idx, table in enumerate(tables, 1):
                    for row_idx, row in enumerate(table, 1):
                        for col_idx, cell in enumerate(row, 1):
                            if cell and str(cell).strip():
                                segments.append({
                                    'content': str(cell).strip(),
                                    'location': f'Page {page_num}, Table {table_idx}, Row {row_idx}, Col {col_idx}',
                                    'location_type': 'table_cell',
                                    'page_number': page_num,
                                    'table_number': table_idx,
                                    'row': row_idx,
                                    'column': col_idx,
                                    'speaker': None,
                                    'context_before': '',
                                    'context_after': '',
                                })
        
        # Add context
        for i, seg in enumerate(segments):
            if i > 0:
                seg['context_before'] = segments[i-1]['content'][:100]
            if i < len(segments) - 1:
                seg['context_after'] = segments[i+1]['content'][:100]
                
    except ImportError:
        return [{'error': 'pdfplumber not installed. Run: pip install pdfplumber'}]
    except Exception as e:
        return [{'error': str(e)}]
    
    return segments


def extract_text_from_xlsx(file_path: str) -> list[dict]:
    """Extract text from Excel files with sheet/row/column tracking."""
    segments = []
    try:
        import openpyxl
        
        wb = openpyxl.load_workbook(file_path, data_only=True)
        
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            headers = []
            
            for row_idx, row in enumerate(sheet.iter_rows(), 1):
                # Capture headers from first row
                if row_idx == 1:
                    headers = [cell.value for cell in row]
                    continue
                
                for col_idx, cell in enumerate(row, 1):
                    if cell.value is not None and str(cell.value).strip():
                        content = str(cell.value).strip()
                        header = headers[col_idx-1] if col_idx <= len(headers) and headers[col_idx-1] else f'Column {col_idx}'
                        
                        segments.append({
                            'content': content,
                            'location': f'Sheet "{sheet_name}", Row {row_idx}, Column "{header}"',
                            'location_type': 'cell',
                            'sheet': sheet_name,
                            'row': row_idx,
                            'column': col_idx,
                            'column_header': header,
                            'speaker': None,
                            'context_before': '',
                            'context_after': '',
                        })
                        
    except ImportError:
        return [{'error': 'openpyxl not installed. Run: pip install openpyxl'}]
    except Exception as e:
        return [{'error': str(e)}]
    
    return segments


def extract_text_from_csv(file_path: str) -> list[dict]:
    """Extract text from CSV files with row/column tracking."""
    segments = []
    try:
        import csv
        
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            # Detect delimiter
            sample = f.read(8192)
            f.seek(0)
            try:
                dialect = csv.Sniffer().sniff(sample)
            except:
                dialect = csv.excel
            
            reader = csv.reader(f, dialect)
            headers = []
            
            for row_idx, row in enumerate(reader, 1):
                if row_idx == 1:
                    headers = row
                    continue
                
                for col_idx, cell in enumerate(row, 1):
                    if cell and cell.strip():
                        header = headers[col_idx-1] if col_idx <= len(headers) else f'Column {col_idx}'
                        
                        segments.append({
                            'content': cell.strip(),
                            'location': f'Row {row_idx}, Column "{header}"',
                            'location_type': 'cell',
                            'row': row_idx,
                            'column': col_idx,
                            'column_header': header,
                            'speaker': None,
                            'context_before': '',
                            'context_after': '',
                        })
                        
    except Exception as e:
        return [{'error': str(e)}]
    
    return segments


def extract_text_from_json(file_path: str) -> list[dict]:
    """Extract text from JSON files with path tracking."""
    segments = []
    
    def extract_recursive(obj, path=''):
        if isinstance(obj, dict):
            for key, value in obj.items():
                new_path = f'{path}.{key}' if path else key
                extract_recursive(value, new_path)
        elif isinstance(obj, list):
            for idx, item in enumerate(obj):
                new_path = f'{path}[{idx}]'
                extract_recursive(item, new_path)
        elif isinstance(obj, str) and obj.strip():
            segments.append({
                'content': obj.strip(),
                'location': f'Path: {path}',
                'location_type': 'json_path',
                'json_path': path,
                'speaker': None,
                'context_before': '',
                'context_after': '',
            })
    
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        extract_recursive(data)
    except Exception as e:
        return [{'error': str(e)}]
    
    return segments


def extract_text_from_html(file_path: str) -> list[dict]:
    """Extract text from HTML files."""
    segments = []
    try:
        from html.parser import HTMLParser
        
        class TextExtractor(HTMLParser):
            def __init__(self):
                super().__init__()
                self.texts = []
                self.current_tag = None
                self.tag_stack = []
                
            def handle_starttag(self, tag, attrs):
                self.tag_stack.append(tag)
                self.current_tag = tag
                
            def handle_endtag(self, tag):
                if self.tag_stack:
                    self.tag_stack.pop()
                self.current_tag = self.tag_stack[-1] if self.tag_stack else None
                
            def handle_data(self, data):
                text = data.strip()
                if text and self.current_tag not in ['script', 'style']:
                    self.texts.append({
                        'content': text,
                        'tag': self.current_tag,
                    })
        
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            content = f.read()
        
        parser = TextExtractor()
        parser.feed(content)
        
        for idx, item in enumerate(parser.texts, 1):
            segments.append({
                'content': item['content'],
                'location': f'Element {idx} ({item["tag"] or "text"})',
                'location_type': 'html_element',
                'element_index': idx,
                'tag': item['tag'],
                'speaker': None,
                'context_before': '',
                'context_after': '',
            })
            
    except Exception as e:
        return [{'error': str(e)}]
    
    return segments


# ============================================================================
# FILE TYPE ROUTER
# ============================================================================

EXTRACTORS = {
    '.txt': extract_text_from_txt,
    '.md': extract_text_from_txt,
    '.docx': extract_text_from_docx,
    '.pdf': extract_text_from_pdf,
    '.xlsx': extract_text_from_xlsx,
    '.xls': extract_text_from_xlsx,
    '.csv': extract_text_from_csv,
    '.tsv': extract_text_from_csv,
    '.json': extract_text_from_json,
    '.html': extract_text_from_html,
    '.htm': extract_text_from_html,
}

SUPPORTED_EXTENSIONS = set(EXTRACTORS.keys())


def extract_text(file_path: str) -> list[dict]:
    """Route to appropriate extractor based on file type."""
    ext = Path(file_path).suffix.lower()
    extractor = EXTRACTORS.get(ext)
    
    if not extractor:
        return [{'error': f'Unsupported file type: {ext}'}]
    
    return extractor(file_path)


# ============================================================================
# SEARCH & MATCHING
# ============================================================================

def tokenize_query(query: str) -> list[str]:
    """
    Break query into search terms.
    Handles phrases in quotes, AND/OR operators, and individual terms.
    """
    terms = []
    
    # Extract quoted phrases first
    quoted = re.findall(r'"([^"]+)"', query)
    terms.extend(quoted)
    
    # Remove quoted phrases from query
    remaining = re.sub(r'"[^"]+"', '', query)
    
    # Remove operators
    remaining = re.sub(r'\b(AND|OR|NOT)\b', ' ', remaining, flags=re.IGNORECASE)
    
    # Split remaining into words, filter out short/common words
    stopwords = {'the', 'a', 'an', 'is', 'are', 'was', 'were', 'be', 'been', 
                 'being', 'have', 'has', 'had', 'do', 'does', 'did', 'will',
                 'would', 'could', 'should', 'may', 'might', 'must', 'shall',
                 'can', 'need', 'to', 'of', 'in', 'for', 'on', 'with', 'at',
                 'by', 'from', 'as', 'into', 'through', 'during', 'before',
                 'after', 'above', 'below', 'between', 'under', 'again',
                 'further', 'then', 'once', 'here', 'there', 'when', 'where',
                 'why', 'how', 'all', 'each', 'few', 'more', 'most', 'other',
                 'some', 'such', 'no', 'nor', 'not', 'only', 'own', 'same',
                 'so', 'than', 'too', 'very', 'just', 'and', 'but', 'if', 'or',
                 'because', 'until', 'while', 'about', 'what', 'which', 'who',
                 'this', 'that', 'these', 'those', 'it', 'its'}
    
    words = remaining.lower().split()
    for word in words:
        word = re.sub(r'[^\w\-]', '', word)  # Keep hyphens for compound terms
        if word and len(word) > 2 and word not in stopwords:
            terms.append(word)
    
    return list(set(terms))  # Deduplicate


def calculate_match_score(content: str, terms: list[str]) -> tuple[float, list[str]]:
    """
    Calculate relevance score and identify which terms matched.
    Returns (score, matched_terms).
    """
    content_lower = content.lower()
    matched = []
    
    for term in terms:
        term_lower = term.lower()
        if term_lower in content_lower:
            matched.append(term)
    
    if not matched:
        return 0.0, []
    
    # Score based on:
    # 1. Percentage of query terms matched
    term_coverage = len(matched) / len(terms)
    
    # 2. Density of matches (matches per 100 words)
    word_count = len(content.split())
    match_count = sum(content_lower.count(t.lower()) for t in matched)
    density = min(1.0, match_count / max(1, word_count / 10))
    
    # 3. Exact phrase bonus (if multi-word terms match exactly)
    phrase_bonus = 0
    for term in matched:
        if ' ' in term and term.lower() in content_lower:
            phrase_bonus += 0.2
    
    score = (term_coverage * 0.5) + (density * 0.3) + min(0.2, phrase_bonus)
    
    return round(score, 3), matched


def search_segments(segments: list[dict], terms: list[str], 
                    min_score: float = 0.0) -> list[dict]:
    """Search extracted segments for matching content."""
    matches = []
    
    for seg in segments:
        if 'error' in seg:
            continue
            
        content = seg.get('content', '')
        if not content:
            continue
        
        score, matched_terms = calculate_match_score(content, terms)
        
        if score > min_score and matched_terms:
            seg['match_score'] = score
            seg['search_terms_matched'] = matched_terms
            matches.append(seg)
    
    # Sort by score descending
    matches.sort(key=lambda x: x.get('match_score', 0), reverse=True)
    
    return matches


# ============================================================================
# DATE DETECTION
# ============================================================================

def detect_date_in_file(file_path: str, segments: list[dict]) -> Optional[str]:
    """
    Try to detect the date of the research from file metadata or content.
    """
    # First try file modification date
    try:
        mtime = os.path.getmtime(file_path)
        file_date = datetime.fromtimestamp(mtime).strftime('%Y-%m-%d')
    except:
        file_date = None
    
    # Look for dates in filename
    filename = Path(file_path).stem
    date_patterns = [
        r'(\d{4}[-_]\d{2}[-_]\d{2})',  # 2024-01-15
        r'(\d{2}[-_]\d{2}[-_]\d{4})',  # 15-01-2024
        r'(\d{4}\d{2}\d{2})',           # 20240115
        r'(Q[1-4][-_\s]?\d{4})',        # Q1 2024
        r'(\w+[-_\s]\d{4})',            # January 2024
    ]
    
    for pattern in date_patterns:
        match = re.search(pattern, filename)
        if match:
            return match.group(1)
    
    # Look for dates in first few segments
    for seg in segments[:10]:
        content = seg.get('content', '')
        for pattern in date_patterns:
            match = re.search(pattern, content)
            if match:
                return match.group(1)
    
    return file_date


# ============================================================================
# MAIN SEARCH ORCHESTRATION  
# ============================================================================

def find_research_files(folders: list[str]) -> list[str]:
    """Recursively find all supported files in given folders."""
    files = []
    
    for folder in folders:
        folder_path = Path(folder)
        if not folder_path.exists():
            print(f"Warning: Folder not found: {folder}", file=sys.stderr)
            continue
        
        if folder_path.is_file():
            if folder_path.suffix.lower() in SUPPORTED_EXTENSIONS:
                files.append(str(folder_path))
            continue
        
        for ext in SUPPORTED_EXTENSIONS:
            files.extend(str(f) for f in folder_path.rglob(f'*{ext}'))
    
    return sorted(set(files))


def analyze_file(file_path: str, search_terms: list[str]) -> FileAnalysis:
    """Analyze a single file for matching content."""
    path = Path(file_path)
    
    analysis = FileAnalysis(
        file_path=str(path.absolute()),
        file_name=path.name,
        file_type=path.suffix.lower(),
        file_size=path.stat().st_size,
        modified_date=datetime.fromtimestamp(path.stat().st_mtime).strftime('%Y-%m-%d %H:%M'),
        evidence_count=0,
    )
    
    # Extract text
    segments = extract_text(file_path)
    
    # Check for extraction errors
    if segments and 'error' in segments[0]:
        analysis.extraction_errors.append(segments[0]['error'])
        return analysis
    
    # Detect date
    detected_date = detect_date_in_file(file_path, segments)
    
    # Search for matches
    matches = search_segments(segments, search_terms)
    
    # Convert to Evidence objects
    for match in matches:
        evidence = Evidence(
            content=match['content'],
            file_path=str(path.absolute()),
            file_name=path.name,
            file_type=path.suffix.lower(),
            location=match.get('location', 'Unknown'),
            location_type=match.get('location_type', 'unknown'),
            context_before=match.get('context_before', '')[:200],
            context_after=match.get('context_after', '')[:200],
            speaker=match.get('speaker'),
            date_detected=detected_date,
            match_score=match.get('match_score', 0),
            search_terms_matched=match.get('search_terms_matched', []),
        )
        analysis.evidence.append(evidence)
    
    analysis.evidence_count = len(analysis.evidence)
    
    return analysis


def run_search(folders: list[str], query: str) -> SearchResults:
    """Execute full search across all folders."""
    search_terms = tokenize_query(query)
    
    if not search_terms:
        print("Error: No valid search terms extracted from query", file=sys.stderr)
        sys.exit(1)
    
    print(f"Search terms: {search_terms}", file=sys.stderr)
    
    results = SearchResults(
        query=query,
        search_terms=search_terms,
        folders_searched=folders,
        files_analyzed=0,
        files_with_matches=0,
        total_evidence_count=0,
        search_timestamp=datetime.now().isoformat(),
    )
    
    # Find all files
    files = find_research_files(folders)
    results.files_analyzed = len(files)
    
    print(f"Found {len(files)} files to analyze", file=sys.stderr)
    
    # Analyze each file
    for i, file_path in enumerate(files, 1):
        print(f"  [{i}/{len(files)}] {Path(file_path).name}...", file=sys.stderr, end='')
        
        try:
            analysis = analyze_file(file_path, search_terms)
            
            if analysis.evidence_count > 0:
                results.files_with_matches += 1
                results.total_evidence_count += analysis.evidence_count
                results.file_analyses.append(analysis)
                print(f" {analysis.evidence_count} matches", file=sys.stderr)
            else:
                print(f" no matches", file=sys.stderr)
                
            if analysis.extraction_errors:
                results.errors.extend(
                    f"{analysis.file_name}: {e}" for e in analysis.extraction_errors
                )
                
        except Exception as e:
            results.errors.append(f"{file_path}: {str(e)}")
            print(f" ERROR: {e}", file=sys.stderr)
    
    # Sort file analyses by total evidence count
    results.file_analyses.sort(key=lambda x: x.evidence_count, reverse=True)
    
    return results


def results_to_dict(results: SearchResults) -> dict:
    """Convert results to JSON-serializable dict."""
    return {
        'query': results.query,
        'search_terms': results.search_terms,
        'folders_searched': results.folders_searched,
        'files_analyzed': results.files_analyzed,
        'files_with_matches': results.files_with_matches,
        'total_evidence_count': results.total_evidence_count,
        'search_timestamp': results.search_timestamp,
        'errors': results.errors,
        'file_analyses': [
            {
                'file_path': fa.file_path,
                'file_name': fa.file_name,
                'file_type': fa.file_type,
                'file_size': fa.file_size,
                'modified_date': fa.modified_date,
                'evidence_count': fa.evidence_count,
                'evidence': [asdict(e) for e in fa.evidence],
                'extraction_errors': fa.extraction_errors,
            }
            for fa in results.file_analyses
        ]
    }


# ============================================================================
# CLI
# ============================================================================

def main():
    parser = argparse.ArgumentParser(
        description='Search research files exhaustively for relevant content.',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python search_files.py ./research --query "onboarding friction"
  python search_files.py ./interviews ./surveys --query "pricing concerns" -o results.json
  python search_files.py . --query "enterprise SSO requirements"
        """
    )
    
    parser.add_argument('folders', nargs='+', help='Folders to search (recursive)')
    parser.add_argument('--query', '-q', required=True, help='Search query')
    parser.add_argument('--output', '-o', help='Output JSON file (default: stdout)')
    parser.add_argument('--min-score', type=float, default=0.0, 
                        help='Minimum match score threshold (0-1)')
    
    args = parser.parse_args()
    
    print(f"\n{'='*60}", file=sys.stderr)
    print(f"Knowledge Finder - Exhaustive Research Search", file=sys.stderr)
    print(f"{'='*60}", file=sys.stderr)
    print(f"Query: {args.query}", file=sys.stderr)
    print(f"Folders: {args.folders}", file=sys.stderr)
    print(f"{'='*60}\n", file=sys.stderr)
    
    results = run_search(args.folders, args.query)
    
    print(f"\n{'='*60}", file=sys.stderr)
    print(f"SEARCH COMPLETE", file=sys.stderr)
    print(f"  Files analyzed: {results.files_analyzed}", file=sys.stderr)
    print(f"  Files with matches: {results.files_with_matches}", file=sys.stderr)
    print(f"  Total evidence pieces: {results.total_evidence_count}", file=sys.stderr)
    if results.errors:
        print(f"  Errors: {len(results.errors)}", file=sys.stderr)
    print(f"{'='*60}\n", file=sys.stderr)
    
    # Output results
    output_data = results_to_dict(results)
    
    if args.output:
        with open(args.output, 'w', encoding='utf-8') as f:
            json.dump(output_data, f, indent=2, ensure_ascii=False)
        print(f"Results saved to: {args.output}", file=sys.stderr)
    else:
        print(json.dumps(output_data, indent=2, ensure_ascii=False))


if __name__ == '__main__':
    main()
